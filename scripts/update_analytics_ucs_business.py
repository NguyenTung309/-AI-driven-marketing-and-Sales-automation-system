from docx import Document

SRC = r"C:\Users\admin\OneDrive\Desktop\New Microsoft Word Document.docx"
doc = Document(SRC)

# Description row = 3, Trigger row = 4

# ── UC-53 ──
doc.tables[0].cell(3, 1).text = (
    "Cho phep PM theo doi hieu suatmarketing tren tat ca kenh ket noi (Facebook, Zalo, Instagram, TikTok, YouTube) "
    "tren mot dashboard duy nhat. Moi ngay he thong tu dong tong hop du lieu tu leads, hoi thoai, tin nhan, "
    "chi phi quang cao va doanh thu thanh cac chi so: so lead, so tin nhan den, so phan hoi tu dong hoa (AI), "
    "ti le chuyen doi, thoi gian phan hoi trung binh, chi phi quang cao, chi phi/lead (CPL) va doanh thu. "
    "PM co the xem theo tung kenh hoac tong hop, so sanh voi ky truoc (ngay-sang-ngay hoac tuan-sang-tuan) "
    "de nhan biet xuan huong. Dashboard tu dong lam moi du lieu moi 60 giay va bao trang thai du lieu cu "
    "neu qua 36 gio chua cap nhat."
)
doc.tables[0].cell(4, 1).text = (
    "PM truy cap trang Bao cao > Tong quan. He thong tu dong tai du lieu KPI da tong hop tu bang kpi_daily "
    "va hien thi tren dashboard. PM co the chon khoang thoi gian (7 ngay hoac 30 ngay) va loc theo kenh cu the."
)

# ── UC-54 ──
doc.tables[1].cell(3, 1).text = (
    "Cho phep PM theo doi luong chuyen doi tu lead thanh khach hang qua 4 buoc: "
    "Lead den -> Tin nhan (hoi thoai) -> Phan hoi (tin nhan tu dong hoa) -> Chuyen doi (khach hang). "
    "Moi buoc hien thi so luong va ty le chuyen tiep so voi buoc truoc. "
    "PM co the loc theo kenh cu the (Facebook, Zalo, Instagram) de xem hieu suat chuyen doi "
    "cua tung kenh rieng biet, hoac xem tat ca kenh cong lai. "
    "Bieu do thanh ngang hien thi gia tri cu the va phan tram o moi buoc de PM de dang danh gia "
    "luong nao dang bi roi nat trong quy trinh."
)
doc.tables[1].cell(4, 1).text = (
    "PM truy cap trang Bao cao > Tong quan hoac Chuyen doi Lead, chon khoang thoi gian va kenh. "
    "He thong tinh toan va hien thi luong chuyen doi tu dong."
)

# ── UC-55 ──
doc.tables[2].cell(3, 1).text = (
    "Cho phep PM danh gia hieu suat cua tung AI Agent trong he thong: "
    "so luot xu ly, ty le hoan thanh, chat luong phan hoi (danh gia qua mau chat luong), "
    "luong su dung token AI (input/output) va chi phi tung agent. "
    "PM co the so sanh giua cac agent de biet agent nao xu ly nhieu nhat, "
    "agent nao co ty le loi cao, va agent nao ton kem nhat de tu dieu chinh cau hinh. "
    "Bang du lieu sap xep theo so luot xu ly giam dan, voi cot chi phi AI tinh theo USD "
    "va cot ty le loi (100% - ty le hoan thanh). "
    "Radar chart hieu ung top 5 agent theo ty le hoan thanh giup PM nhin tong the."
)
doc.tables[2].cell(4, 1).text = (
    "PM truy cap trang Bao cao > Hieu suat Agent (yeu co quyen analytics.tab.agent). "
    "He thong tai du lieu hieu suat agent va chi phi AI cho khoang thoi gian da chon."
)

# ── UC-56 ──
doc.tables[3].cell(3, 1).text = (
    "He thong tu dong thuc hien tong hop KPI hang ngay luc 00:30 GMT+7 cho moi tenant. "
    "Quy trinh doc du lieu thô tu cac bang leads, conversations, messages, ads_campaigns, "
    "ads_metrics_dailies va lead_revenues, sau do tinh toan va ghi vao bang kpi_daily. "
    "Moi kenh (Facebook, Zalo, Instagram, TikTok, YouTube) duoc tach rieng va co dong tong hop 'all'. "
    "Chi so duoc tinh gom: so lead, so tin nhan den (DM), so phan hoi tu dong, so hoi thoai co phan hoi, "
    "so chuyen doi, thoi gian phan hoi trung binh, tong chi phi quang cao va doanh thu duoc duyet. "
    "Day la nguon du lieu chinh cho tat ca cac bao cao va dashboard tren he thong."
)
doc.tables[3].cell(4, 1).text = (
    "He thong tu dong kich hoat theo lich (cron) luc 00:30 GMT+7 moi ngay cho tung tenant. "
    "Khong can su tuong tac cua nguoi dung."
)

# ── UC-57 ──
doc.tables[4].cell(3, 1).text = (
    "Cho phep PM xuat du lieu KPI da tong hop duoi dang tep CSV hoac PDF de chia se voi doi ngu. "
    "Du lieu xuat giong chinh xac voi du lieu hien thi tren dashboard omnichannel: "
    "moi kenh voi cac chi so lead, tin nhan, phan hoi, chuyen doi, thoi gian phan hoi, "
    "chi phi quang cao va CPL. "
    "Dinh dang CSV phu hop de mo trong Excel hoac cong cu phan tich khac. "
    "Dinh dang PDF dung QuestPDF voi dinh dang A4, font Calibri, bang 8 cot "
    "phu hop de in hoac gui noi bo. "
    "Ten tep tu dong tao theo format analytics-{tu ngay}-{den ngay}.{dinh dang}."
)
doc.tables[4].cell(4, 1).text = (
    "PM bam nut 'Xuat bao cao' tren header trang Bao cao, chon dinh dang CSV hoac PDF tu hop thoai xuat."
)

# ── UC-58 ──
doc.tables[5].cell(3, 1).text = (
    "He thong tu dong phat hien bat thuong tren cac chi so KPI su dung phuong phap phan tich "
    "thong ke z-score tren chuoi thoi gian lich su. Khi gia tri cua mot chi so (vi du CPL, so lead, "
    "chi phi quang cao) vuot qua nguong z-score mac dinh (|z| > 3), he thong danh dau la bat thuong. "
    "PM co the xem danh sach bat thuong tren tab Chuyen doi Lead, voi thong tin: "
    "kenh nao, chi so gi, ngay nao, gia tri bao nhieu, z-score bao nhieu, va co phai bat thuong khong. "
    "Muc dich giup PM nhanh chong nhan biet CPL tang dot ngot, lead giam bat thuong "
    "hoac thoi gian phan hoi lech so voi trung binh 14 ngay gan day de dieu chinh ke hoach."
)
doc.tables[5].cell(4, 1).text = (
    "PM truy cap trang Bao cao > Chuyen doi Lead. He thong tu dong phan tich "
    "va hien thi danh sach bat thuong cho chi so CPL (mac dinh) voi cua so 14 ngay."
)

# ── UC-71 ──
doc.tables[6].cell(3, 1).text = (
    "Cho phep PM/Sales Lead du bao luong lead cho 7 ngay toi dua tren du lieu lich su. "
    "He thong su dung mo hinh phan tich chuoi thoi gian (ML.NET TimeSeries SSA) "
    "da duoc huan luyen tren it nhat 60 ngay du lieu lich su de tao du bao. "
    "Moi diem du bao bao gom: gia tri du bao, gioi tren va gioi duoi (confidence bounds) "
    "giup PM hieu khoang dao dong co the xay ra. "
    "Du bao duoc cap nhat moi 24 gio va hien thi tren bieu do SVG voi truc X la ngay "
    "va truc Y la so luong lead du bao. "
    "PM co the chon kenh cu the de xem du bao rieng cho tung nen tang."
)
doc.tables[6].cell(4, 1).text = (
    "PM/Sales Lead truy cap trang Bao cao > Chuyen doi Lead. "
    "He thong tu dong tai du bao lead 7 ngay cho kenh tat ca hoac kenh da chon."
)

doc.save(SRC)
print("Done - Descriptions and Triggers updated.")
