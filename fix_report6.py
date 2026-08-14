import docx
from docx.shared import Pt

DST = r"E:\DoAn\-AI-driven-marketing-and-Sales-automation-system\Report6_Software_User_Guides.docx"
doc = docx.Document(DST)

# Fix 1: Remove placeholder text in Deliverable Package section (P27)
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("[The section will list"):
        for run in para.runs:
            run.clear()
        para.clear()
        run = para.add_run("The following table lists all deliverable items included in this release:")
        run.font.size = Pt(11)
        break

# Fix 2: Remove leftover placeholder from template workflow (P115)
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("[Describe the detailed guides"):
        for run in para.runs:
            run.clear()
        para.clear()
        # Leave empty or add useful text
        break

# Fix 3: Verify deliverable table has descriptions
pkg_table = doc.tables[1]
for ri in range(1, len(pkg_table.rows)):
    cell = pkg_table.rows[ri].cells[2]
    if not cell.text.strip():
        cell.text = "See project documentation"

doc.save(DST)
print("Fixes applied successfully")
print(f"Paragraphs: {len(doc.paragraphs)}")
