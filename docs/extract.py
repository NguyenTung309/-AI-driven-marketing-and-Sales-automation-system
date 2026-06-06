import sys
from docx import Document
d = Document(sys.argv[1])
out = []
for p in d.paragraphs:
    t = p.text.strip()
    if t:
        out.append(t)
for t in d.tables:
    for row in t.rows:
        cells = [c.text.strip().replace('\n',' ') for c in row.cells]
        out.append(' | '.join(cells))
sys.stdout.reconfigure(encoding='utf-8')
print('\n'.join(out))
