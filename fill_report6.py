import copy
import os
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\admin\Downloads\SEP490 Templates-New\Report6_Software User Guides.docx"
DST = r"E:\DoAn\-AI-driven-marketing-and-Sales-automation-system\Report6_Software_User_Guides.docx"

shutil.copy2(SRC, DST)
doc = Document(DST)

# Helper: clear a paragraph and write styled text
def set_para_text(para, text, bold=False, italic=False, size=None, color=None, font_name=None, align=None):
    for run in para.runs:
        run.clear()
    para.clear()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if align is not None:
        para.alignment = align
    return run

def add_para_after(para, text, bold=False, italic=False, size=None, color=None, font_name=None, align=None):
    new_para = OxmlElement('w:p')
    para._element.addnext(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, para._parent)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if font_name:
        run.font.name = font_name
    if align is not None:
        p.alignment = align
    return p

def set_cell_text(cell, text, bold=False, size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ========== 1. Record of Changes Table ==========
changes_table = doc.tables[0]
# Fill header row (row 0 already has headers)
# Add rows with data
changes_data = [
    ("01/08/2026", "A", "Dev Team", "Initial Software User Guides document created"),
    ("01/08/2026", "A", "Dev Team", "Added system requirements and installation instructions"),
    ("01/08/2026", "A", "Dev Team", "Added user manual for login, dashboard, inbox, sale assist, leads, documents, content, agents, and admin workflows"),
]

for i, (date, action, person, desc) in enumerate(changes_data):
    row_idx = i + 1
    if row_idx >= len(changes_table.rows):
        row = changes_table.add_row()
    else:
        row = changes_table.rows[row_idx]
    set_cell_text(row.cells[0], date, size=10)
    set_cell_text(row.cells[1], action, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row.cells[2], person, size=10)
    set_cell_text(row.cells[3], desc, size=10)

# ========== 2. Deliverable Package Table ==========
pkg_table = doc.tables[1]
deliverables = [
    ("1", "Project Schedule/Tracking", "Project plan and sprint tracking documents stored in docs/ClawBot_SaleMkt_ProjectPlan.docx"),
    ("2", "Project Backlog", "Product backlog managed via SDD (Spec-Driven Development) artifacts in .sdd/ directory"),
    ("3", "Source Codes", "Full source code: 12 .NET projects (Clawbot.sln), React 19 frontend (clawbot-web), gRPC proto contracts, deployment scripts"),
    ("4", "Database Script(s)", "DDL migration scripts in deploy/migrations/ (0001_init.sql + incremental SQL files), SQL Server 2022"),
    ("5", "Final Report Document", "Capstone project reports (Report 1 through Report 6)"),
    ("6", "Test Cases Document", "xUnit test suites: Domain tests, Application tests, Integration tests (Testcontainers), and frontend component tests"),
    ("7", "Defects List", "Tracked via Git issues and SDD defect tracking in .sdd/ artifacts"),
    ("8", "Issues List", "Technical debt and known issues documented in docs/module-checklist.md and docs/spec-audit.md"),
    ("9", "Slide", "Presentation slide deck for capstone project defense"),
]

for i, (no, item, desc) in enumerate(deliverables):
    row_idx = i + 1
    if row_idx < len(pkg_table.rows):
        row = pkg_table.rows[row_idx]
        set_cell_text(row.cells[2], desc, size=10)

# ========== 3. System Requirements ==========
# Find the "2.1 System Requirements" section and fill content
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "2.1 System Requirements" in para.text:
        # The next paragraph is the placeholder
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()
            run = target.add_run("The ClawBot SaleMkt system requires the following software and hardware configurations:")
            run.font.size = Pt(11)
            
            # Add sub-items after this paragraph
            reqs = [
                "Server-Side Requirements:",
                "  - .NET SDK 8.0.x (ASP.NET Core 8 runtime)",
                "  - SQL Server 2022 (database engine)",
                "  - Redis 7.x (caching and SignalR backplane)",
                "  - RabbitMQ (message broker via MassTransit)",
                "  - Qdrant (vector database for Knowledge Base RAG)",
                "  - MinIO (S3-compatible object storage)",
                "  - Docker Desktop 4.x+ (for infrastructure containers)",
                "",
                "Client-Side Requirements:",
                "  - Modern web browser (Chrome 90+, Firefox 88+, Edge 90+, Safari 14+)",
                "  - Stable internet connection (minimum 5 Mbps recommended)",
                "",
                "Development Environment:",
                "  - Node.js 20+ and npm (for frontend build)",
                "  - Visual Studio 2022 with ASP.NET/.NET workload (recommended)",
                "  - PowerShell or Windows Terminal",
                "  - Git for version control",
            ]
            
            last_p = target
            for req in reqs:
                new_p = add_para_after(last_p, req, bold=req.endswith(":"), size=11)
                last_p = new_p
        break

# ========== 4. Installation Instructions ==========
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "2.2 Installation Instruction" in para.text:
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()
            run = target.add_run("Follow the steps below to install and configure ClawBot SaleMkt on a local development machine:")
            run.font.size = Pt(11)
            
            steps = [
                ("Step 1: Clone the repository", "Clone the source code from the Git repository to your local machine using git clone command."),
                ("Step 2: Install prerequisites", "Install .NET SDK 8.0, Node.js 20+, Docker Desktop, and verify installation using dotnet --version, node --version, and docker info commands."),
                ("Step 3: Configure environment", "Copy deploy/.env.example to deploy/.env. Configure environment variables including MSSQL_SA_PASSWORD (default: Clawbot!2026), RABBITMQ_USER/PASSWORD (default: guest/guest), MINIO_USER/PASSWORD (default: minio/minio12345), and optional LLM provider keys."),
                ("Step 4: Start infrastructure containers", "Run docker compose --env-file deploy/.env -f deploy/docker-compose.yml up -d to start SQL Server, Redis, RabbitMQ, Qdrant, MinIO, and Metabase containers."),
                ("Step 5: Initialize database", "Create the clawbot database using SQL Server container and apply migration scripts from deploy/migrations/ directory. Alternatively, run run-all.bat which handles database setup automatically."),
                ("Step 6: Build the solution", "Execute dotnet restore Clawbot.sln followed by dotnet build Clawbot.sln --no-restore to compile all 12 .NET projects."),
                ("Step 7: Start backend services", "Open separate terminal windows and start AgentService (port 15875), API backend (port 15874), and Gateway (port 15873) with ASPNETCORE_ENVIRONMENT set to Development."),
                ("Step 8: Start frontend", "Navigate to src/frontend/clawbot-web, run npm ci to install dependencies, then npm run dev -- --host 0.0.0.0 --port 15876 to start the React development server."),
                ("Step 9: Access the application", "Open http://localhost:15876 in your browser. Log in using the default admin account: email admin@clawbot.local, password Admin@12345."),
            ]
            
            last_p = target
            for step_title, step_desc in steps:
                p1 = add_para_after(last_p, "", size=11)
                p2 = add_para_after(p1, step_title, bold=True, size=11)
                p3 = add_para_after(p2, step_desc, size=11)
                last_p = p3
        break

# ========== 5. User Manual - Overview ==========
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "3.1 Overview" in para.text:
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()
            run = target.add_run("ClawBot SaleMkt is an AI-driven omnichannel marketing and sales automation system designed for Chinese language education centers. The system integrates customer conversations from multiple channels (Zalo, Facebook, TikTok, Instagram, YouTube) into a unified inbox, provides AI-powered sales assistance, automated lead scoring and pipeline management, document generation, content management, and comprehensive analytics dashboards.")
            run.font.size = Pt(11)
            
            overview_items = [
                "",
                "Key features of the system include:",
                "  - Omnichannel Inbox: Unified conversation management from 5 social channels via Pancake integration",
                "  - Sale Assist: AI-powered draft replies, conversation summaries, and quick reply templates",
                "  - Lead CRM: Automated lead scoring, pipeline stages (Hot/Warm/Cold), and assignment",
                "  - Document Automation: PDF quote generation, brochures, onboarding kits with tenant branding",
                "  - Content Pipeline: Content brief creation, approval workflow, scheduling, and publishing",
                "  - Agent Dashboard: 8 AI agents with monitoring, configuration, and trace capabilities",
                "  - Knowledge Base: Chinese language module management with versioning and accuracy testing",
                "  - Analytics: Omnichannel KPIs, funnel analysis, agent performance, and anomaly detection",
                "  - Admin Console: User/role management, API keys, integrations, and audit logs",
                "",
                "The system architecture follows Clean Architecture with Domain-Driven Design (DDD) bounded contexts, using .NET 8 backend, React 19 frontend, gRPC agent communication, and SQL Server 2022 database.",
            ]
            
            last_p = target
            for item in overview_items:
                if item:
                    bold = item.endswith(":")
                    new_p = add_para_after(last_p, item, bold=bold, size=11)
                    last_p = new_p
                else:
                    last_p = add_para_after(last_p, "", size=6)
        break

# ========== 6. Workflow 1: Login and Authentication ==========
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "3.2 Workflow 1" in para.text:
        # Update heading text
        for run in para.runs:
            run.clear()
        run = para.add_run("3.2 Workflow 1: Login and Authentication")
        run.bold = True
        
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()
            run = target.add_run("Purpose: This workflow describes how users authenticate into the ClawBot system, including login, two-factor authentication (2FA), and password recovery.")
            run.font.size = Pt(11)
            
            workflow1_items = [
                "",
                "Step-by-step guide:",
                "",
                "1. Navigate to http://localhost:15876/login in your web browser.",
                "2. Enter your registered email address and password in the login form.",
                "3. Click the Login button to submit credentials.",
                "4. If 2FA is enabled, enter the 6-digit OTP code from your authenticator app.",
                "5. Upon successful authentication, you will be redirected to the Dashboard (/).",
                "6. The system issues a JWT access token valid for 60 minutes.",
                "",
                "Password Recovery:",
                "  1. Click Forgot Password on the login page.",
                "  2. Enter your registered email address.",
                "  3. Check your email for a 6-digit OTP code.",
                "  4. Enter the OTP and set a new password.",
                "  5. Click Confirm to complete the reset process.",
                "",
                "Note: The system locks accounts after 5 failed login attempts for 15 minutes. Contact your administrator if your account is locked.",
            ]
            
            last_p = target
            for item in workflow1_items:
                if item:
                    bold = item.endswith(":") or item.startswith("Step-by-step")
                    new_p = add_para_after(last_p, item, bold=bold, size=11)
                    last_p = new_p
                else:
                    last_p = add_para_after(last_p, "", size=6)
        break

# ========== 7. Workflow 2: Dashboard and Analytics ==========
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "3.3 Workflow 2" in para.text:
        for run in para.runs:
            run.clear()
        run = para.add_run("3.3 Workflow 2: Dashboard Overview and Analytics")
        run.bold = True
        
        if i + 1 < len(doc.paragraphs):
            target = doc.paragraphs[i + 1]
            target.clear()
            run = target.add_run("Purpose: This workflow demonstrates how administrators and managers monitor system performance through the Dashboard and Analytics modules.")
            run.font.size = Pt(11)
            
            workflow2_items = [
                "",
                "Step-by-step guide:",
                "",
                "1. After login, the Dashboard (/) displays at the home page.",
                "2. View the overview metrics: total leads, messages received, response rate, conversion rate, and AI agent cost.",
                "3. Review the omnichannel chart showing message volume by platform (Zalo, Facebook, TikTok, etc.).",
                "4. Check the funnel analysis for lead conversion stages.",
                "5. Navigate to Analytics (/analytics) for detailed reports.",
                "6. Select a date range to filter analytics data.",
                "7. Review three report tabs: Overview (omnichannel KPIs), Agent (performance per agent), and Lead (scoring and conversion).",
                "8. Export reports as CSV or PDF if needed.",
                "9. Check the Notifications center (/notifications) for system alerts, hot lead notifications, and anomaly warnings.",
            ]
            
            last_p = target
            for item in workflow2_items:
                if item:
                    bold = item.endswith(":") or item.startswith("Step-by-step")
                    new_p = add_para_after(last_p, item, bold=bold, size=11)
                    last_p = new_p
                else:
                    last_p = add_para_after(last_p, "", size=6)
        break

# ========== 8. Add more workflows after Workflow 2 ==========
# Find the last content paragraph after Workflow 2 and add more workflows
last_workflow_para = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "3.3 Workflow 2" in para.text:
        # Find the next heading 1 or end of content
        for j in range(i + 2, len(doc.paragraphs)):
            p = doc.paragraphs[j]
            if p.style.name == "Heading 1" or p.style.name == "Heading 2":
                break
            last_workflow_para = p
        break

# We need to insert new headings and content. Let me find the right insertion point.
# The template only has up to 3.3. We need to add 3.4, 3.5, 3.6, 3.7, 3.8, 3.9
# Find the position after the last workflow 2 content

insert_after = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name == "Heading 3" and "3.3 Workflow 2" in para.text:
        # Find the last paragraph in this section
        for j in range(i + 2, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name in ("Heading 1", "Heading 2", "Heading 3"):
                insert_after = doc.paragraphs[j - 1]
                break
        break

if insert_after is None:
    # Fallback: find last paragraph
    insert_after = doc.paragraphs[-1]

# Additional workflows to add
additional_workflows = [
    ("3.4 Workflow 3: Omnichannel Inbox Management", [
        "Purpose: This workflow describes how sales staff manage customer conversations from multiple channels in a unified inbox.",
        "",
        "Step-by-step guide:",
        "",
        "1. Navigate to Conversations (/conversations) from the sidebar.",
        "2. The left panel displays a list of all conversations with status indicators.",
        "3. Filter conversations by status (All, Open, Escalated, Mine) or by channel.",
        "4. Use the search bar to find conversations by customer name, phone number, or ID.",
        "5. Click on a conversation to open it in the center chat panel.",
        "6. Read the message history including inbound (customer) and outbound (AI/sale) messages.",
        "7. Type a reply in the composer. Toggle the AI assistant on/off as needed.",
        "8. Click Send to deliver the response through the original channel.",
        "9. Review the right context panel for customer information, lead score, and suggested actions.",
        "10. Use Escalate or Resolve buttons to change conversation status.",
        "",
        "The system uses SignalR for real-time updates, so new messages appear automatically without page refresh.",
    ]),
    ("3.5 Workflow 4: AI Sale Assist", [
        "Purpose: This workflow demonstrates how AI assists sales staff in drafting responses, summarizing conversations, and providing quick replies.",
        "",
        "Step-by-step guide:",
        "",
        "1. Open a conversation in the Inbox (/conversations).",
        "2. In the conversation view, locate the Sale Assist panel on the right side.",
        "3. Click Create Draft to generate an AI-suggested response.",
        "4. The system uses RAG (Retrieval-Augmented Generation) with Knowledge Base data and Claude AI to generate a contextual reply.",
        "5. Review the drafted response in the chat panel.",
        "6. Click Approve and Send to use the draft, or Discard to reject it.",
        "7. For long conversations, click Summary to get a 3-bullet AI summary of the conversation.",
        "8. Use Quick Reply templates for common responses (e.g., pricing inquiry, trial class invitation).",
        "9. The system also suggests next actions based on conversation context.",
        "",
        "Note: The AI assistant does not replace sales staff. It reduces reading and drafting time while the sales person makes the final decision on what to send.",
    ]),
    ("3.6 Workflow 5: Lead Management and CRM", [
        "Purpose: This workflow covers lead creation, scoring, pipeline management, and assignment in the CRM module.",
        "",
        "Step-by-step guide:",
        "",
        "1. Navigate to Leads (/leads) from the sidebar.",
        "2. View the summary cards: total leads, hot leads, unassigned, and 7-day forecast.",
        "3. Switch between Table view and Kanban board view (5 columns: Hot, Warm, Cold, Customer, Lost).",
        "4. In Table view, filter leads by name, source, status, or assigned owner.",
        "5. Click on a lead card or row to open the Lead Drawer (slide-in panel).",
        "6. Review the Timeline tab for lead activity history.",
        "7. Check the Context tab for contact information and suggested follow-up actions.",
        "8. View the Revenue tab for revenue tracking (approve/reject).",
        "9. Create a new lead manually or let the system auto-create from widget/omnichannel interactions.",
        "10. The system automatically scores leads based on configured rules (e.g., replied to message = +40 points, platform = Zalo).",
        "11. Lead stage updates automatically: score >= 70 = Hot, 30-70 = Warm, < 30 = Cold.",
        "12. Use the Assign action to manually reassign leads to specific sales staff.",
    ]),
    ("3.7 Workflow 6: Document Generation", [
        "Purpose: This workflow shows how to generate sales documents such as quotes, brochures, and onboarding kits.",
        "",
        "Step-by-step guide:",
        "",
        "1. Navigate to Documents (/documents) from the sidebar.",
        "2. Browse the template library showing available document templates.",
        "3. Select a template (e.g., Quote, HSK Brochure, Onboarding Kit).",
        "4. Fill in the required fields in the form (customer name, course details, pricing).",
        "5. Click Generate to create the document. The system merges template data with customer information.",
        "6. Preview the generated document before downloading.",
        "7. Click Download to save the PDF file.",
        "8. Use Generate Kit to create multiple documents at once (quote + brochure + onboarding kit).",
        "",
        "Note: Document generation requires configuration of document storage (MinIO) and email delivery (SMTP) for live sending. The system tracks who created each document, for which customer, and when it was sent.",
    ]),
    ("3.8 Workflow 7: Knowledge Base Management", [
        "Purpose: This workflow describes managing the Chinese language Knowledge Base with modules, versions, and accuracy testing.",
        "",
        "Step-by-step guide:",
        "",
        "1. Navigate to Knowledge Base (/kb) from the sidebar.",
        "2. The left Module Rail displays available KB modules (e.g., HSK Curriculum, Course Paths, Pricing, FAQ).",
        "3. Click on a module to view its version history in the center Version Rail.",
        "4. Select a version to view or edit its content in the right Editor Workspace.",
        "5. Create a new version by clicking New Version and editing the content.",
        "6. Deploy a version by clicking Deploy. The system embeds content into Qdrant vector store for RAG retrieval.",
        "7. Rollback to a previous version if needed using the Rollback button.",
        "8. Use QA Test Cases to validate KB accuracy. Add test questions and expected answers.",
        "9. Click Run Test to execute the test set against the deployed version.",
        "10. Review accuracy results. The system alerts if accuracy drops below 85%.",
    ]),
    ("3.9 Workflow 8: Admin System Configuration", [
        "Purpose: This workflow covers system administration tasks including user management, roles, API keys, and integration configuration.",
        "",
        "Step-by-step guide:",
        "",
        "1. Navigate to System (/system) from the sidebar footer.",
        "2. Users tab: Create new users, edit profiles, reset passwords, enable/disable accounts.",
        "3. Roles tab: Create roles (Admin, Sale, Marketer, QA, Viewer), assign permissions to roles.",
        "4. API Keys tab: Generate, revoke, and rotate API keys for external integrations.",
        "5. Integrations tab: Configure Pancake channel connection (base URL, access token, webhook secret).",
        "6. System Logs tab: View system-level logs with cursor-based pagination.",
        "7. Audit tab: Review audit trail of administrative actions.",
        "8. Navigate to Profile (/profile) to manage your personal information.",
        "9. In Profile, change your password or enable/disable two-factor authentication.",
        "10. Configure LLM providers at /llm-providers to manage AI model settings and API keys.",
    ]),
]

# Insert additional workflows
current_element = insert_after._element
for wf_title, wf_lines in additional_workflows:
    # Add heading
    heading_p = OxmlElement('w:p')
    heading_pPr = OxmlElement('w:pPr')
    heading_style = OxmlElement('w:pStyle')
    heading_style.set(qn('w:val'), 'Heading3')
    heading_pPr.append(heading_style)
    heading_p.append(heading_pPr)
    heading_r = OxmlElement('w:r')
    heading_rPr = OxmlElement('w:rPr')
    heading_b = OxmlElement('w:b')
    heading_rPr.append(heading_b)
    heading_r.append(heading_rPr)
    heading_t = OxmlElement('w:t')
    heading_t.text = wf_title
    heading_t.set(qn('xml:space'), 'preserve')
    heading_r.append(heading_t)
    heading_p.append(heading_r)
    current_element.addnext(heading_p)
    current_element = heading_p
    
    # Add content lines
    for line in wf_lines:
        content_p = OxmlElement('w:p')
        content_r = OxmlElement('w:r')
        content_rPr = OxmlElement('w:rPr')
        content_sz = OxmlElement('w:sz')
        content_sz.set(qn('w:val'), '22')  # 11pt
        content_rPr.append(content_sz)
        content_r.append(content_rPr)
        content_t = OxmlElement('w:t')
        content_t.text = line
        content_t.set(qn('xml:space'), 'preserve')
        content_r.append(content_t)
        content_p.append(content_r)
        current_element.addnext(content_p)
        current_element = content_p
    
    # Add spacing paragraph
    spacing_p = OxmlElement('w:p')
    spacing_r = OxmlElement('w:r')
    spacing_t = OxmlElement('w:t')
    spacing_t.text = ""
    spacing_r.append(spacing_t)
    spacing_p.append(spacing_r)
    current_element.addnext(spacing_p)
    current_element = spacing_p

doc.save(DST)
print(f"Document saved to {DST}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
