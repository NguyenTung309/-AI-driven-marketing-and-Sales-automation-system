"""
Convert USER-FLOWS.md to a professional DOCX using compact_reference_guide preset.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT = "Calibri"
FONT_MONO = "Consolas"
BODY_SIZE = Pt(11)
BODY_AFTER = Pt(6)
BODY_LINE = 1.25
H1_SIZE = Pt(16)
H1_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H1_BEFORE = Pt(18)
H1_AFTER = Pt(10)
H2_SIZE = Pt(13)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H2_BEFORE = Pt(14)
H2_AFTER = Pt(7)
H3_SIZE = Pt(12)
H3_COLOR = RGBColor(0x1F, 0x4D, 0x78)
H3_BEFORE = Pt(10)
H3_AFTER = Pt(5)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER_COLOR = "BFBFBF"
MARGIN = Inches(1.0)
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, font_name=FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)


def set_para_spacing(para, before=Pt(0), after=BODY_AFTER, line=BODY_LINE):
    fmt = para.paragraph_format
    fmt.space_before = before
    fmt.space_after = after
    fmt.line_spacing = line


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = BODY_AFTER
    style.paragraph_format.line_spacing = BODY_LINE
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), FONT)
        rFonts.set(qn("w:hAnsi"), FONT)

    for level, (sz, clr, bef, aft) in {
        1: (H1_SIZE, H1_COLOR, H1_BEFORE, H1_AFTER),
        2: (H2_SIZE, H2_COLOR, H2_BEFORE, H2_AFTER),
        3: (H3_SIZE, H3_COLOR, H3_BEFORE, H3_AFTER),
    }.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT
        hs.font.size = sz
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = bef
        hs.paragraph_format.space_after = aft
        hs.paragraph_format.line_spacing = 1.2


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def parse_md_tables(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r"\|[\s\-:|]+\|", line):
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
    return rows


def add_table(doc, rows, col_widths=None):
    if not rows:
        return None
    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    if col_widths is None:
        col_w = CONTENT_WIDTH_DXA // n_cols
        col_widths = [col_w] * n_cols

    tbl_pr = tbl._element.find(qn("w:tblPr"))
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{CONTENT_WIDTH_DXA}" w:type="dxa"/>')
    old = tbl_pr.find(qn("w:tblW"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(tblW)

    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="120" w:type="dxa"/>')
    old = tbl_pr.find(qn("w:tblInd"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(tblInd)

    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
    </w:tblBorders>'''
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(parse_xml(borders_xml))

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j in range(n_cols):
            cell = row.cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            set_para_spacing(para, before=Pt(2), after=Pt(2), line=1.15)
            cell_text = row_data[j] if j < len(row_data) else ""
            if "**" in cell_text:
                parts = re.split(r"(\*\*.*?\*\*)", cell_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = para.add_run(part[2:-2])
                        set_run_font(run, bold=True, size=Pt(9))
                    else:
                        run = para.add_run(part)
                        set_run_font(run, size=Pt(9))
            else:
                run = para.add_run(cell_text)
                set_run_font(run, size=Pt(9))

            tc = cell._element
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
                tc.insert(0, tcPr)
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[j]}" w:type="dxa"/>')
            old = tcPr.find(qn("w:tcW"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcW)

            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            old = tcPr.find(qn("w:vAlign"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(vAlign)

            tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
                <w:top w:w="60" w:type="dxa"/>
                <w:bottom w:w="60" w:type="dxa"/>
                <w:start w:w="100" w:type="dxa"/>
                <w:end w:w="100" w:type="dxa"/>
            </w:tcMar>''')
            old = tcPr.find(qn("w:tcMar"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcMar)

        if i == 0:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
                    tc.insert(0, tcPr)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_FILL}" w:val="clear"/>')
                old = tcPr.find(qn("w:shd"))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(shading)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    return tbl


def add_rich_text(para, text, default_size=BODY_SIZE, default_bold=False):
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            set_run_font(run, font_name=FONT_MONO, size=Pt(9))
        else:
            run = para.add_run(part)
            set_run_font(run, size=default_size, bold=default_bold)


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(6), after=Pt(6), line=1.0)
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        para._element.insert(0, pPr)
    pBdr = parse_xml(f'''<w:pBdr {nsdecls("w")}>
        <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D0D0"/>
    </w:pBdr>''')
    pPr.append(pBdr)


def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    para = doc.add_paragraph()
                    set_para_spacing(para, before=Pt(4), after=Pt(4), line=1.0)
                    pPr = para._element.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
                        para._element.insert(0, pPr)
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                    pPr.append(shd)
                    pBdr = parse_xml(f'''<w:pBdr {nsdecls("w")}>
                        <w:top w:val="single" w:sz="4" w:space="1" w:color="D0D0D0"/>
                        <w:left w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>
                        <w:bottom w:val="single" w:sz="4" w:space="1" w:color="D0D0D0"/>
                        <w:right w:val="single" w:sz="4" w:space="4" w:color="D0D0D0"/>
                    </w:pBdr>''')
                    pPr.append(pBdr)
                    run = para.add_run(code_text)
                    set_run_font(run, font_name=FONT_MONO, size=Pt(8))
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Flush pending table
        if in_table and not stripped.startswith("|"):
            if table_lines:
                rows = parse_md_tables(table_lines)
                if rows:
                    add_table(doc, rows)
                table_lines = []
            in_table = False

        # Table row
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text_clean = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
            para = doc.add_heading(text_clean, level=min(level, 3))
            for run in para.runs:
                set_run_font(run, font_name=FONT,
                           size={1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}.get(level, H3_SIZE),
                           bold=True,
                           color={1: H1_COLOR, 2: H2_COLOR, 3: H3_COLOR}.get(level, H3_COLOR))
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.375)
            para.paragraph_format.first_line_indent = Inches(-0.2)
            set_para_spacing(para, before=Pt(1), after=Pt(3))
            run = para.add_run(f"{num}. ")
            set_run_font(run, bold=True, color=H2_COLOR)
            add_rich_text(para, text)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.375)
            para.paragraph_format.first_line_indent = Inches(-0.2)
            set_para_spacing(para, before=Pt(1), after=Pt(3))
            run = para.add_run("\u2022  ")
            set_run_font(run, color=H2_COLOR)
            add_rich_text(para, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            set_para_spacing(para, before=Pt(2), after=Pt(2))
            pPr = para._element.find(qn("w:pPr"))
            if pPr is None:
                pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
                para._element.insert(0, pPr)
            pBdr = parse_xml(f'''<w:pBdr {nsdecls("w")}>
                <w:left w:val="single" w:sz="12" w:space="6" w:color="2E74B5"/>
+            </w:pBdr>''')
            pPr.append(pBdr)
            add_rich_text(para, text, default_size=Pt(10))
            i += 1
            continue

        # Regular paragraph
        para = doc.add_paragraph()
        set_para_spacing(para)
        add_rich_text(para, stripped)
        i += 1

    # Flush remaining table
    if in_table and table_lines:
        rows = parse_md_tables(table_lines)
        if rows:
            add_table(doc, rows)

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    md_path = Path(r"E:\DoAn\-AI-driven-marketing-and-Sales-automation-system\docs\screen-flows\USER-FLOWS.md")
    docx_path = md_path.with_suffix(".docx")
    result = convert_md_to_docx(md_path, docx_path)
    print(f"Created: {result}")
