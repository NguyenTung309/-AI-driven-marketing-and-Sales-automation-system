"""Fix SITEMAP-UPDATED.docx: all 6 columns visible."""
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT = "Calibri"
FONT_MONO = "Consolas"

def cell_xml(tag, attrs, text_attrs=""):
    return f"<w:{tag} {nsdecls('w')}{attrs}{text_attrs}/>"

def convert(tsv_path, docx_path):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Inches(8.5)
    s.page_height = Inches(11)
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.6)
    s.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    rows = []
    with open(tsv_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f, delimiter="\t")
        for row in reader:
            if any(c.strip() for c in row):
                rows.append(row)

    # Title
    p = doc.add_paragraph()
    r = p.add_run("Updated Sitemap")
    r.font.name = FONT
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Version 1.0.0 | 2026-08-06 | Draft")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(8)

    # Table: 7.3in usable (0.6in margins), columns sum to 9600 dxa = 6.67in
    # Plenty of room for all content
    col_w = [720, 1900, 1900, 2200, 1500, 880]
    TOTAL_W = sum(col_w)
    n_rows = len(rows)
    n_cols = len(col_w)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.autofit = False

    tbl_pr = tbl._element.find(qn("w:tblPr"))
    # tblW
    old = tbl_pr.find(qn("w:tblW"))
    if old is not None:
        tbl_pr.remove(old)
    el = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{TOTAL_W}" w:type="dxa"/>')
    tbl_pr.append(el)
    # tblInd
    old = tbl_pr.find(qn("w:tblInd"))
    if old is not None:
        tbl_pr.remove(old)
    el = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tbl_pr.append(el)
    # borders
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    bxml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(parse_xml(bxml))

    HEADER_FILL = "E8EEF5"
    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        is_hdr = (i == 0)
        for j in range(n_cols):
            cell = row.cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.0

            cell_text = row_data[j] if j < len(row_data) else ""
            run = para.add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(8)

            if is_hdr:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
            elif j == 0:
                run.font.name = FONT_MONO
                run.bold = True
            elif j == 2:
                run.font.name = FONT_MONO
                run.font.size = Pt(7.5)
            elif j == 5:
                run.bold = True
                if "PRIMARY" in cell_text:
                    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
                elif "DEEP" in cell_text:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                else:
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            tc = cell._element
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
                tc.insert(0, tcPr)

            # Cell width
            old = tcPr.find(qn("w:tcW"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_w[j]}" w:type="dxa"/>'))

            # Vertical align
            old = tcPr.find(qn("w:vAlign"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))

            # Cell margins
            old = tcPr.find(qn("w:tcMar"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                '<w:top w:w="30" w:type="dxa"/>'
                '<w:bottom w:w="30" w:type="dxa"/>'
                '<w:start w:w="60" w:type="dxa"/>'
                '<w:end w:w="60" w:type="dxa"/>'
                '</w:tcMar>'
            ))

            # Shading
            old = tcPr.find(qn("w:shd"))
            if old is not None:
                tcPr.remove(old)
            if is_hdr:
                fill = HEADER_FILL
            elif i % 2 == 0:
                fill = "F7F9FC"
            else:
                fill = "FFFFFF"
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>'))

    # Footer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("PRIMARY = main page | DEEP = deep-link/route | SUB = sub-component/tab | PUBLIC = no auth")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(str(docx_path))
    return docx_path

if __name__ == "__main__":
    tsv = Path(r"E:\DoAn\-AI-driven-marketing-and-Sales-automation-system\docs\screen-flows\SITEMAP-UPDATED.tsv")
    out = tsv.with_suffix(".docx")
    print("Created:", convert(tsv, out))
