"""Convert SITEMAP-UPDATED.tsv to DOCX."""
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT = "Calibri"
FONT_MONO = "Consolas"
BODY_SIZE = Pt(11)
BODY_AFTER = Pt(6)
BODY_LINE = 1.25
H1_SIZE = Pt(16)
H1_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H1_BEFORE = Pt(18)
H1_AFTER = Pt(10)
H2_SIZE = Pt(13)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H2_BEFORE = Pt(14)
H2_AFTER = Pt(7)
H3_SIZE = Pt(12)
H3_COLOR = RGBColor(0x1F, 0x4D, 0x78)
H3_BEFORE = Pt(10)
H3_AFTER = Pt(5)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER_COLOR = "BFBFBF"
MARGIN = Inches(1.0)
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, font_name=FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)


def set_para_spacing(para, before=Pt(0), after=BODY_AFTER, line=BODY_LINE):
    fmt = para.paragraph_format
    fmt.space_before = before
    fmt.space_after = after
    fmt.line_spacing = line


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = BODY_AFTER
    style.paragraph_format.line_spacing = BODY_LINE
    for level, (sz, clr, bef, aft) in {
        1: (H1_SIZE, H1_COLOR, H1_BEFORE, H1_AFTER),
        2: (H2_SIZE, H2_COLOR, H2_BEFORE, H2_AFTER),
        3: (H3_SIZE, H3_COLOR, H3_BEFORE, H3_AFTER),
    }.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT
        hs.font.size = sz
        hs.font.color.rgb = clr
        hs.font.bold = True
        hs.paragraph_format.space_before = bef
        hs.paragraph_format.space_after = aft
        hs.paragraph_format.line_spacing = 1.2


def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_table(doc, rows, col_widths=None, compact_font_size=Pt(8.5)):
    if not rows:
        return
    n_cols = len(rows[0])
    n_rows = len(rows)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    usable = Inches(7.0)
    usable_dxa = int(usable / 914400 * 1440 * 640)
    if col_widths is None:
        cw = usable_dxa // n_cols
        col_widths = [cw] * n_cols

    tbl_pr = tbl._element.find(qn("w:tblPr"))
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{usable_dxa}" w:type="dxa"/>')
    old = tbl_pr.find(qn("w:tblW"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(tblW)

    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    old = tbl_pr.find(qn("w:tblInd"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(tblInd)

    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>
    </w:tblBorders>'''
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    tbl_pr.append(parse_xml(borders_xml))

    # Level color map for visual grouping
    level_colors = {
        "PRIMARY": RGBColor(0x1F, 0x4D, 0x78),
        "PRIMARY under section": RGBColor(0x2E, 0x74, 0xB5),
        "DEEP": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-02a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-03a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from notification": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-05a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-06a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-07a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-07b1": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-09a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-09a1": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-09b": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-09e": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-09a (routes)": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-10a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-11a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-11b": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-11c": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-12a": RGBColor(0x4A, 0x4A, 0x4A),
        "DEEP - from P-13c": RGBColor(0x4A, 0x4A, 0x4A),
        "SUB": RGBColor(0x66, 0x66, 0x66),
    }

    def get_level_color(level_text):
        if level_text in level_colors:
            return level_colors[level_text]
        if "DEEP" in level_text:
            return RGBColor(0x4A, 0x4A, 0x4A)
        if "PRIMARY" in level_text:
            return RGBColor(0x1F, 0x4D, 0x78)
        if "SUB" in level_text:
            return RGBColor(0x66, 0x66, 0x66)
        return RGBColor(0x33, 0x33, 0x33)

    # Alternate row shading
    ROW_SHADE_NORMAL = "FFFFFF"
    ROW_SHADE_ALT = "F7F9FC"

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        is_header = (i == 0)
        for j in range(n_cols):
            cell = row.cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            set_para_spacing(para, before=Pt(1), after=Pt(1), line=1.1)
            cell_text = row_data[j] if j < len(row_data) else ""

            run = para.add_run(cell_text)
            if is_header:
                set_run_font(run, bold=True, size=Pt(8.5))
                run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
            elif j == 5:  # Level column
                set_run_font(run, size=Pt(8), bold=True, color=get_level_color(cell_text))
            elif j == 0:  # P-ID column - mono
                set_run_font(run, font_name=FONT_MONO, size=Pt(8.5), bold=True)
            elif j == 2:  # Route column - mono
                set_run_font(run, font_name=FONT_MONO, size=Pt(8))
            else:
                set_run_font(run, size=Pt(8.5))

            tc = cell._element
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
                tc.insert(0, tcPr)

            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_widths[j]}" w:type="dxa"/>')
            old = tcPr.find(qn("w:tcW"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcW)

            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            old = tcPr.find(qn("w:vAlign"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(vAlign)

            tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
                <w:top w:w="40" w:type="dxa"/>
                <w:bottom w:w="40" w:type="dxa"/>
                <w:start w:w="80" w:type="dxa"/>
                <w:end w:w="80" w:type="dxa"/>
            </w:tcMar>''')
            old = tcPr.find(qn("w:tcMar"))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcMar)

        # Header row
        if is_header:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
                    tc.insert(0, tcPr)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_FILL}" w:val="clear"/>')
                old = tcPr.find(qn("w:shd"))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(shading)
        else:
            # Alternate row shading
            shade = ROW_SHADE_ALT if (i % 2 == 0) else ROW_SHADE_NORMAL
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = parse_xml(f"<w:tcPr {nsdecls('w')}/>")
                    tc.insert(0, tcPr)
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade}" w:val="clear"/>')
                old = tcPr.find(qn("w:shd"))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(shading)


def add_rich_text(para, text, default_size=BODY_SIZE):
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            set_run_font(run, font_name=FONT_MONO, size=Pt(9))
        else:
            run = para.add_run(part)
            set_run_font(run, size=default_size)


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(6), after=Pt(6), line=1.0)
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        para._element.insert(0, pPr)
    pBdr = parse_xml(f'''<w:pBdr {nsdecls("w")}>
        <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D0D0"/>
    </w:pBdr>''')
    pPr.append(pBdr)


import re


def convert_tsv_to_docx(tsv_path, docx_path):
    doc = Document()
    setup_styles(doc)
    setup_page(doc)

    # Title
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(0), after=Pt(4))
    run = para.add_run("Updated Sitemap")
    set_run_font(run, size=Pt(22), bold=True, color=H1_COLOR)

    # Subtitle metadata
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(0), after=Pt(2))
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = parse_xml(f"<w:pPr {nsdecls('w')}/>")
        para._element.insert(0, pPr)
    pBdr = parse_xml(f'''<w:pBdr {nsdecls("w")}>
        <w:left w:val="single" w:sz="12" w:space="6" w:color="2E74B5"/>
    </w:pBdr>''')
    pPr.append(pBdr)
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run("Version: 1.0.0  |  Date: 2026-08-06  |  Status: Draft")
    set_run_font(run, size=Pt(10), italic=True, color=RGBColor(0x66, 0x66, 0x66))

    add_horizontal_rule(doc)

    # Read TSV
    rows = []
    with open(tsv_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f, delimiter="\t")
        header = next(reader)
        rows.append(header)
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append(row)

    # Summary paragraph
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(6), after=Pt(10))
    run = para.add_run(f"Sitemap contains {len(rows)} pages across all levels: ")
    set_run_font(run, size=Pt(10))

    # Count by level
    level_counts = {}
    for row in rows:
        if len(row) >= 6:
            lv = row[5].strip()
            # Normalize level categories
            if "PRIMARY" in lv:
                key = "PRIMARY"
            elif "DEEP" in lv:
                key = "DEEP"
            elif "SUB" in lv:
                key = "SUB"
            else:
                key = lv
            level_counts[key] = level_counts.get(key, 0) + 1

    summary_parts = [f"{v} {k}" for k, v in sorted(level_counts.items())]
    run = para.add_run(", ".join(summary_parts))
    set_run_font(run, size=Pt(10), bold=True)

    # Build the main sitemap table
    # Column widths (usable ~7.0in = ~10080 DXA at 0.75in margins)
    # P-ID: ~0.7in, Page Name: ~1.6in, Route: ~1.5in, Component: ~1.6in, Roles: ~1.2in, Level: ~0.6in
    col_widths = [900, 2100, 1900, 2100, 1600, 1000]

    add_table(doc, rows, col_widths=col_widths)

    # Footer note
    add_horizontal_rule(doc)
    para = doc.add_paragraph()
    set_para_spacing(para, before=Pt(4), after=Pt(0))
    run = para.add_run("Legend: ")
    set_run_font(run, size=Pt(9), bold=True)
    run = para.add_run("PRIMARY = main page  |  ")
    set_run_font(run, size=Pt(9))
    run = para.add_run("DEEP = deep-link / route  |  ")
    set_run_font(run, size=Pt(9))
    run = para.add_run("SUB = sub-component / tab  |  ")
    set_run_font(run, size=Pt(9))
    run = para.add_run("PUBLIC = no auth required")
    set_run_font(run, size=Pt(9))

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    tsv_path = Path(r"E:\DoAn\-AI-driven-marketing-and-Sales-automation-system\docs\screen-flows\SITEMAP-UPDATED.tsv")
    docx_path = tsv_path.with_suffix(".docx")
    result = convert_tsv_to_docx(tsv_path, docx_path)
    print(f"Created: {result}")
